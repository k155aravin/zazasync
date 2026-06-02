from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\Users\Fly Pirate\Documents\New project\zazasync\ZazaSync_Team_Product_Brief.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 31, 36)
MUTED = RGBColor(92, 102, 112)
LIGHT = "F2F4F7"
SOFT_GREEN = "EAF7E8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def style_run(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        style_run(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run = p.add_run(item)
        style_run(run, size=10.8, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        style_run(run, size=10.8, color=INK)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        style_run(run, size=16 if level == 1 else 13, color=BLUE if level == 1 else DARK_BLUE, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_callout(doc, label, body, fill=SOFT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="BFD8BA")
    cell = table.cell(0, 0)
    set_width(cell, 6.5)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    style_run(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    style_run(r2, size=11, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_width(cell, widths[i])
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        style_run(r, size=10.2, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_width(cell, widths[i])
            set_cell_margins(cell, top=90, bottom=90)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            style_run(r, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    return doc


def build():
    doc = setup_doc()
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.text = "ZazaSync Team Brief"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(header.runs[0], size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.text = "Prepared for internal product discussion"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(footer.runs[0], size=9, color=MUTED)

    add_para(doc, "ZazaSync", size=28, color=INK, bold=True, after=2)
    add_para(doc, "Team Product Brief: Simple UX, Useful Features, Clear Direction", size=13, color=MUTED, after=14)
    add_para(doc, "Prepared June 1, 2026", size=10, color=MUTED, after=18)

    add_callout(
        doc,
        "Core idea",
        "ZazaSync should help adults 21+ find an SQDC product, see where it appears available, and get alerted when it comes back in stock.",
    )

    add_heading(doc, "What We Are Building")
    add_para(
        doc,
        "ZazaSync is a Quebec SQDC product intelligence platform. It should be cleaner and more product-focused than a raw inventory crawler: users search for the product first, then check availability, freshness, and alerts.",
    )
    add_bullets(
        doc,
        [
            "Product-first discovery: search by product, brand, category, format, price, or potency.",
            "Availability evidence: show where the product appears available and when it was checked.",
            "Back-in-stock alerts: if a product is out of stock, let the user set an email or SMS alert.",
            "Simple profile value: saved products, preferred stores, and alert settings.",
        ],
    )

    add_heading(doc, "Best User Experience")
    add_callout(
        doc,
        "UX principle",
        "The site should answer one practical question: Can I buy this today, and where?",
        fill="F4F6F9",
    )
    add_numbered(
        doc,
        [
            "User lands on the homepage and immediately sees a large product search bar.",
            "User searches for a product, brand, category, or phrase like 'pre-roll under $20'.",
            "Results show simple product cards with name, brand, format, price, THC/CBD, and stock status.",
            "User opens a product and sees nearby SQDC availability, last checked time, and related products.",
            "If unavailable, user taps 'Alert me when available' and chooses email, SMS, or both.",
            "Profile appears only when useful: saved products, preferred stores, and notification settings.",
        ],
    )

    add_heading(doc, "Recommended Site Structure")
    add_table(
        doc,
        ["Section", "Purpose", "Must Answer"],
        [
            ("Inventory", "Browse and search all tracked products.", "What products match what I want?"),
            ("Product Detail", "One product, all important facts.", "Where is it available and when was it checked?"),
            ("New Drops", "Show newly detected products.", "What is new at SQDC?"),
            ("Back in Stock", "Show recently returned products.", "What came back recently?"),
            ("Stores", "Show tracked SQDC locations.", "Which stores have useful availability evidence?"),
            ("Watchlist", "Personal saved products and alerts.", "What am I waiting for or tracking?"),
        ],
        [1.25, 2.45, 2.8],
    )

    add_heading(doc, "Features WeedCrawler Does Not Emphasize")
    add_table(
        doc,
        ["Feature", "Why It Matters", "MVP Priority"],
        [
            ("Back-in-stock alerts", "Gives users a reason to return and create an account.", "High"),
            ("Preferred stores", "Makes availability personal instead of showing every location equally.", "High"),
            ("Saved products", "Creates a simple watchlist and repeat-use loop.", "High"),
            ("Price drop alerts", "Adds value beyond stock status.", "Medium"),
            ("Best value score", "Helps users compare price, format, potency, and availability.", "Medium"),
            ("Recently restocked", "Turns inventory changes into a useful browsing surface.", "High"),
        ],
        [1.55, 3.7, 1.25],
    )

    add_heading(doc, "Homepage Direction")
    add_para(doc, "The homepage should feel like a tool, not a technical report. The first screen should prioritize action.")
    add_bullets(
        doc,
        [
            "Headline: Find SQDC products near you.",
            "Primary control: large search bar.",
            "Quick filters: In stock, New drops, Back in stock, Under $25, Pre-rolls, CBD.",
            "Featured rows: Popular right now, Recently restocked, New drops, Stores with most availability.",
            "Small trust signal: last snapshot time and number of tracked products/stores.",
        ],
    )

    add_heading(doc, "MVP Build Order")
    add_numbered(
        doc,
        [
            "Clean homepage messaging and make Inventory the primary call to action.",
            "Improve product cards so they are simple, readable, and action-oriented.",
            "Create product detail pages with availability evidence and last checked time.",
            "Add 'Alert me' for out-of-stock products with email first.",
            "Add preferred stores and saved products.",
            "Add SMS alerts only after email alerts are working well.",
        ],
    )

    add_heading(doc, "Team Decisions")
    add_bullets(
        doc,
        [
            "Should alerts require an account, or should email-only alerts work without signup?",
            "Should the first release be English-only, French-first, or bilingual?",
            "Should SMS be part of launch, or wait until email alerts prove demand?",
            "What exact product fields are reliable enough to show publicly?",
            "What privacy language is needed before collecting profile and preference data?",
        ],
    )

    add_heading(doc, "Recommendation")
    add_callout(
        doc,
        "Recommended path",
        "Build ZazaSync as a mobile-first web app before a native app. Keep the experience simple: search, availability, alerts, saved products, preferred stores.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
